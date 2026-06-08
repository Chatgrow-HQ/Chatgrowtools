from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
README_PATH = ROOT / "README.md"
OUTPUT_PATH = ROOT / "docs" / "ChatGrow_Free_Tools_Technical_and_Operations_Manual.docx"

BLACK = RGBColor(0, 0, 0)
WHITE = "FFFFFF"
BLACK_HEX = "000000"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.color.rgb = BLACK
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), BLACK_HEX)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def set_repeat_table_cell_text(cell, text: str, *, bold=False, size=9.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, 18, 10),
        "Heading 2": (13, 14, 7),
        "Heading 3": (12, 10, 5),
    }
    for style_name, (size, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.0

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.font.color.rgb = BLACK
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("CHATGROW FREE TOOLS  |  TECHNICAL AND OPERATIONS MANUAL")
    set_run_font(run, size=8.5, bold=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    add_page_number(p)


INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|`.+?`|\[.+?\]\(.+?\))")


def add_inline_markdown(paragraph, text: str, *, default_size=11) -> None:
    position = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run, size=default_size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=default_size, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Courier New", size=default_size - 1)
        else:
            link_match = re.match(r"\[(.+?)\]\((.+?)\)", token)
            if link_match:
                label, url = link_match.groups()
                run = paragraph.add_run(f"{label} ({url})")
                set_run_font(run, size=default_size)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=default_size)


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    add_inline_markdown(paragraph, text)


def add_list_item(doc: Document, text: str, ordered: bool) -> None:
    paragraph = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
    paragraph.paragraph_format.keep_together = True
    add_inline_markdown(paragraph, text)


def add_code_block(doc: Document, lines: list[str]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.right_indent = Inches(0.2)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.keep_together = True
    border = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "5")
        node.set(qn("w:color"), BLACK_HEX)
        border.append(node)
    paragraph._p.get_or_add_pPr().append(border)
    run = paragraph.add_run("\n".join(lines))
    set_run_font(run, name="Courier New", size=8.5)


def table_widths(headers: list[str]) -> list[int]:
    count = len(headers)
    if count == 2:
        first = headers[0].lower()
        if first in {"area", "variable", "format", "layer"}:
            return [2700, 6660]
        return [3600, 5760]
    if count == 3:
        return [2200, 2200, 4960]
    if count == 4:
        return [1400, 1800, 2600, 3560]
    base = CONTENT_WIDTH_DXA // count
    widths = [base] * count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = len(rows[0])
    widths = table_widths(rows[0])
    table = doc.add_table(rows=len(rows), cols=column_count)
    set_table_geometry(table, widths)
    set_table_borders(table)
    repeat_table_header(table.rows[0])

    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        prevent_row_split(row)
        for column_index in range(column_count):
            value = values[column_index] if column_index < len(values) else ""
            set_repeat_table_cell_text(
                row.cells[column_index],
                value,
                bold=row_index == 0,
                size=9 if row_index == 0 else 9.5,
            )

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    run = kicker.add_run("TECHNICAL, OPERATIONAL, ONBOARDING, AND HANDOVER GUIDE")
    set_run_font(run, size=10, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("ChatGrow Free Tools")
    set_run_font(run, size=28, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(34)
    run = subtitle.add_run("Technical and Operations Manual")
    set_run_font(run, size=16)

    description = doc.add_paragraph()
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    description.paragraph_format.left_indent = Inches(0.65)
    description.paragraph_format.right_indent = Inches(0.65)
    description.paragraph_format.space_after = Pt(34)
    run = description.add_run(
        "A complete reference for product walkthroughs, engineering onboarding, "
        "deployment, operations, privacy review, and sale or ownership transfer."
    )
    set_run_font(run, size=11)

    control = doc.add_table(rows=4, cols=2)
    set_table_geometry(control, [2200, 7160])
    set_table_borders(control)
    values = [
        ("Document status", "Current-state manual"),
        ("Documentation date", "June 8, 2026"),
        ("Application", "ChatGrow Free Tools"),
        ("Source of truth", "Repository README.md and application code"),
    ]
    for row, (label, value) in zip(control.rows, values):
        prevent_row_split(row)
        set_repeat_table_cell_text(row.cells[0], label, bold=True)
        set_repeat_table_cell_text(row.cells[1], value)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(22)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run(
        "Black-and-white edition. Verify credentials, domains, model access, "
        "privacy terms, and ownership records at the time of handover."
    )
    set_run_font(run, size=9, italic=True)
    doc.add_page_break()


def add_contents(doc: Document, headings: list[tuple[int, str]]) -> None:
    heading = doc.add_paragraph("Contents", style="Heading 1")
    heading.paragraph_format.space_before = Pt(0)

    for level, text in headings:
        if level != 2:
            continue
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.15)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(text)
        set_run_font(run, size=10.5)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    run = note.add_run(
        "Use Word's Navigation Pane to move through the heading structure."
    )
    set_run_font(run, size=9.5, italic=True)
    doc.add_page_break()


def extract_headings(lines: list[str]) -> list[tuple[int, str]]:
    headings = []
    for line in lines:
        match = re.match(r"^(#{2,3})\s+(.+)$", line)
        if match and match.group(2) != "Table of Contents":
            headings.append((len(match.group(1)), match.group(2)))
    return headings


def content_lines(readme_text: str) -> list[str]:
    lines = readme_text.splitlines()
    start = next(
        index for index, line in enumerate(lines) if line.strip() == "## Product Overview"
    )
    return lines[start:]


def render_markdown(doc: Document, lines: list[str]) -> None:
    index = 0
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            add_body_paragraph(doc, " ".join(part.strip() for part in paragraph_buffer))
            paragraph_buffer = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            index += 1
            continue

        if stripped.startswith("```"):
            flush_paragraph()
            index += 1
            code_lines: list[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            add_code_block(doc, code_lines)
            index += 1
            continue

        if stripped.startswith("|"):
            flush_paragraph()
            rows, index = parse_table(lines, index)
            add_markdown_table(doc, rows)
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            level = min(len(heading_match.group(1)) - 1, 3)
            paragraph = doc.add_paragraph(heading_match.group(2), style=f"Heading {level}")
            if level == 1 and paragraph.text in {
                "System Architecture",
                "Local Setup",
                "Deployment",
                "Privacy and Data Handling",
                "Current Limitations",
                "New Team Member Onboarding",
                "Sale and Technical Handover Checklist",
                "Operational Runbook",
            }:
                paragraph.paragraph_format.page_break_before = True
            index += 1
            continue

        ordered_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if ordered_match:
            flush_paragraph()
            add_list_item(doc, ordered_match.group(1), ordered=True)
            index += 1
            continue

        bullet_match = re.match(r"^-\s+(.+)$", stripped)
        if bullet_match:
            flush_paragraph()
            add_list_item(doc, bullet_match.group(1), ordered=False)
            index += 1
            continue

        paragraph_buffer.append(stripped)
        index += 1

    flush_paragraph()


def add_document_note(doc: Document) -> None:
    doc.add_page_break()
    paragraph = doc.add_paragraph("Document Maintenance", style="Heading 1")
    paragraph.paragraph_format.space_before = Pt(0)
    add_body_paragraph(
        doc,
        "This Word manual is generated from README.md by "
        "`scripts/build_chatgrow_documentation.py`. Update the README first, "
        "then rebuild and visually verify this DOCX whenever product behavior, "
        "tool URLs, deployment, ownership, privacy, or operational procedures change.",
    )


def set_core_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "ChatGrow Free Tools - Technical and Operations Manual"
    props.subject = "Product, engineering, onboarding, operations, and handover documentation"
    props.author = "ChatGrow"
    props.keywords = "ChatGrow, free tools, onboarding, operations, handover, documentation"
    props.comments = "Black-and-white current-state manual generated from the repository README."


def build_document() -> None:
    readme_text = README_PATH.read_text(encoding="utf-8")
    lines = content_lines(readme_text)
    headings = extract_headings(lines)

    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    set_core_properties(doc)
    add_cover(doc)
    add_contents(doc, headings)
    render_markdown(doc, lines)
    add_document_note(doc)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
